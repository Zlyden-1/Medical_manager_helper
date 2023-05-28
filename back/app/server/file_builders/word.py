import copy
import os
from typing import List
from datetime import date
from docx.enum.section import WD_SECTION, WD_ORIENT
from docx.shared import Inches
from docx.shared import RGBColor

import docx


async def builder(filename: str, data: List[dict]) -> str:
    doc = docx.Document()

    doc.sections[0].orientation = WD_ORIENT.LANDSCAPE
    doc.sections[0].page_width = Inches(11.69)
    doc.sections[0].page_height = Inches(8.27)

    table = doc.add_table(rows=1, cols=len(data[0]), )
    heading_cells = table.rows[0].cells
    table.style = 'Table Grid'
    for column in table.columns:
        column.width = Inches(2)

    for i, key in enumerate(data[0].keys()):
        if key != 'Оценка' and key != 'Лишние назначения':
            heading_cells[i].text = key
            heading_cells[i].paragraphs[0].alignment = 1
    for item in data:
        row_cells = table.add_row().cells
        for i, key in enumerate(item.keys()):
            if item["Оценка"] == 'Избыточные назначения':
                if key != 'Оценка' and key != 'Лишние назначения':
                    if type(item[key]) == list:
                        row_cells[i].text = '\n'.join(item[key])
                        row_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 0, 0)
                    else:
                        row_cells[i].text = str(item[key])
                        row_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 0, 0)
                elif key == 'Лишние назначения':
                    cell = table.cell(data.index(item)+1, 7)
                    nested_table = cell.add_table(rows=1, cols=1)
                    nested_cell = nested_table.cell(0, 0)
                    nested_cell.text = '\n'.join(item[key])
                    nested_cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 0, 0)
                    nested_cell.paragraphs[0].runs[0].bold = True
                else:
                    ...
            else:
                if key != 'Оценка':
                    if type(item[key]) == list:
                        row_cells[i].text = '\n'.join(item[key])
                    else:
                        row_cells[i].text = str(item[key])

    # удаляем ячейки в четвертом и пятом столбцах
    for row in table.rows:
        row.cells[8]._element.clear()  # четвертый столбец
        row.cells[9]._element.clear()  # пятый столбец

    # # удаляем четвертый и пятый столбцы
    for row in table.rows:
        row._element.remove(row._element.tc_lst[8])  # четвертый столбец
        row._element.remove(row._element.tc_lst[8])  # пятый столбец

    filename = find_free_filename(filename)
    doc.save(f'../media/{filename}')
    return f'{filename}'


def find_free_filename(base_filename: str):
    filename = f'{base_filename}.docx'
    i = 1
    while os.path.exists(f'../media/{filename}'):
        filename = f"{base_filename}_{i}.docx"
        i += 1
    return filename


if __name__ == '__main__':
    import asyncio

    print(asyncio.run(builder('test', [
                {
                    'Пол пациента': 'Муж',
                    'Дата рождения пациента': date(2000, 1, 1),
                    'ID пациента': 3,
                    'Код МКБ-10': 'J32.9',
                    "Диагноз": "Вазомоторный ринит",
                    'Дата оказания услуги': date(2000, 1, 1),
                    'Должность': 'врач-оториноларинголог',
                    "Назначения": ["Флюорография легких", "Электрокардиография в покое"],
                    "Лишние назначения": ["Креатинин"],
                    "Оценка": "Избыточные назначения",
                    "Источник данных": "Ортоларингологическое_отделение.xlsx",
                    "Дата загрузки": date(2000, 1, 1)
                },
                {
                    'Пол пациента': 'Муж',
                    'Дата рождения пациента': date(2000, 1, 1),
                    'ID пациента': 3,
                    'Код МКБ-10': 'J32.9',
                    "Диагноз": "Вазомоторный ринит",
                    'Дата оказания услуги': date(2000, 1, 1),
                    'Должность': 'врач-оториноларинголог',
                    "Назначения": ["Флюорография легких", "Электрокардиография в покое"],
                    "Лишние назначения": ['ХУЙ'],
                    "Оценка": "Избыточные назначения",
                    "Источник данных": "Ортоларингологическое_отделение.xlsx",
                    "Дата загрузки": date(2000, 2, 1),
                }
            ])))
